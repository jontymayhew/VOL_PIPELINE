"""
scripts/build_docx.py
---------------------
Convert reports/validation_<date>.md → reports/validation_<date>.docx

Steps
-----
1. Style a reference.docx (fonts, colours, margins, header/footer, table style).
2. Run pandoc with that reference doc and the latest validation markdown.

Usage
-----
    python3 scripts/build_docx.py
    python3 scripts/build_docx.py --md reports/validation_20260315.md
"""
from __future__ import annotations

import argparse
import subprocess
import sys
from pathlib import Path

# Colours (hex RGB ints) -------------------------------------------------------
NAVY   = 0x1F3864   # dark navy — headings
MIDBLUE= 0x2E74B5   # mid-blue  — H2
STEEL  = 0x2F5496   # steel     — H3
WHITE  = 0xFFFFFF
LIGHT  = 0xD6E4F0   # table header fill
ALT    = 0xEBF3FB   # alternating table row fill

ROOT = Path(__file__).resolve().parent.parent


# ---------------------------------------------------------------------------
# Style helpers (python-docx)
# ---------------------------------------------------------------------------

def _hex_to_rgb_str(colour_int: int) -> str:
    return f"{colour_int:06X}"


def _set_run_font(run, name: str, size_pt: float,
                  bold: bool = False, color_int: int | None = None):
    from docx.shared import Pt, RGBColor
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color_int is not None:
        r = (color_int >> 16) & 0xFF
        g = (color_int >> 8)  & 0xFF
        b =  color_int        & 0xFF
        run.font.color.rgb = RGBColor(r, g, b)


def _style_paragraph(para, space_before_pt: float = 0, space_after_pt: float = 0,
                      line_spacing_pt: float | None = None):
    from docx.shared import Pt
    para.paragraph_format.space_before = Pt(space_before_pt)
    para.paragraph_format.space_after  = Pt(space_after_pt)
    if line_spacing_pt:
        para.paragraph_format.line_spacing = Pt(line_spacing_pt)


def _set_cell_shading(cell, colour_int: int):
    """Apply solid background fill to a table cell via direct XML."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  _hex_to_rgb_str(colour_int))
    tcPr.append(shd)


def _heading_colour(style, colour_int: int,
                    font_name: str = "Calibri Light",
                    size_pt: float = 14, bold: bool = True):
    from docx.shared import Pt, RGBColor
    f = style.font
    f.name  = font_name
    f.size  = Pt(size_pt)
    f.bold  = bold
    r = (colour_int >> 16) & 0xFF
    g = (colour_int >> 8)  & 0xFF
    b =  colour_int        & 0xFF
    f.color.rgb = RGBColor(r, g, b)


def _get_style(doc, name: str):
    """python-docx __getitem__ has a bug with space-containing names; iterate."""
    return next((s for s in doc.styles if s.name == name), None)


def build_reference_docx(base_path: Path, out_path: Path) -> Path:
    """
    Open the pandoc-default reference.docx, apply professional styles, and save.
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document(str(base_path))

    # -- Page margins --
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.8)
        section.right_margin  = Cm(2.8)

    # -- Default body style --
    body = _get_style(doc, "Normal")
    body.font.name = "Calibri"
    body.font.size = Pt(11)
    body.paragraph_format.space_after  = Pt(6)
    body.paragraph_format.line_spacing = Pt(14)

    # -- Heading styles --
    h1 = _get_style(doc, "Heading 1")
    _heading_colour(h1, NAVY,    size_pt=22, bold=True)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after  = Pt(6)
    h1.paragraph_format.keep_with_next = True

    h2 = _get_style(doc, "Heading 2")
    _heading_colour(h2, MIDBLUE, size_pt=16, bold=True)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after  = Pt(4)
    h2.paragraph_format.keep_with_next = True

    h3 = _get_style(doc, "Heading 3")
    _heading_colour(h3, STEEL,   size_pt=13, bold=True)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after  = Pt(3)
    h3.paragraph_format.keep_with_next = True

    h4 = _get_style(doc, "Heading 4")
    _heading_colour(h4, STEEL,   size_pt=12, bold=True)
    h4.paragraph_format.space_before = Pt(8)
    h4.paragraph_format.space_after  = Pt(2)
    h4.paragraph_format.keep_with_next = True

    # -- Table style: "Table Grid" is the one pandoc uses by default --
    tbl_style = _get_style(doc, "Table Grid")
    if tbl_style:
        tbl_style.font.name = "Calibri"
        tbl_style.font.size = Pt(10)

    # -- Block Text (used by pandoc for blockquotes) --
    bq = _get_style(doc, "Block Text")
    if bq:
        bq.font.name  = "Calibri"
        bq.font.size  = Pt(10)
        bq.font.italic = True
        r = (STEEL >> 16) & 0xFF
        g = (STEEL >> 8)  & 0xFF
        b =  STEEL        & 0xFF
        bq.font.color.rgb = RGBColor(r, g, b)

    # -- Header: report title + separator line --
    for section in doc.sections:
        header = section.header
        header.is_linked_to_previous = False
        if header.paragraphs:
            hp = header.paragraphs[0]
        else:
            hp = header.add_paragraph()
        hp.clear()
        run = hp.add_run("SABR ↔ Local Vol Consistency Validation Report")
        run.font.name  = "Calibri"
        run.font.size  = Pt(9)
        run.font.bold  = True
        r = (NAVY >> 16) & 0xFF
        g = (NAVY >> 8)  & 0xFF
        b =  NAVY        & 0xFF
        run.font.color.rgb = RGBColor(r, g, b)
        hp.paragraph_format.space_after = Pt(2)
        # Bottom border on header paragraph
        pPr = hp._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"),   "single")
        bottom.set(qn("w:sz"),    "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), _hex_to_rgb_str(MIDBLUE))
        pBdr.append(bottom)
        pPr.append(pBdr)

    # -- Footer: page numbers --
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        if footer.paragraphs:
            fp = footer.paragraphs[0]
        else:
            fp = footer.add_paragraph()
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.space_before = Pt(2)
        run = fp.add_run("Page ")
        run.font.name = "Calibri"
        run.font.size = Pt(9)
        # PAGE field
        fld = OxmlElement("w:fldChar")
        fld.set(qn("w:fldCharType"), "begin")
        run._r.append(fld)
        ins = OxmlElement("w:instrText")
        ins.text = "PAGE"
        run._r.append(ins)
        fld2 = OxmlElement("w:fldChar")
        fld2.set(qn("w:fldCharType"), "end")
        run._r.append(fld2)
        run2 = fp.add_run(" of ")
        run2.font.name = "Calibri"
        run2.font.size = Pt(9)
        # NUMPAGES field
        fld3 = OxmlElement("w:fldChar")
        fld3.set(qn("w:fldCharType"), "begin")
        run2._r.append(fld3)
        ins2 = OxmlElement("w:instrText")
        ins2.text = "NUMPAGES"
        run2._r.append(ins2)
        fld4 = OxmlElement("w:fldChar")
        fld4.set(qn("w:fldCharType"), "end")
        run2._r.append(fld4)

    doc.save(str(out_path))
    print(f"  Reference template saved → {out_path}")
    return out_path


# ---------------------------------------------------------------------------
# Post-process: colour table headers and align images
# ---------------------------------------------------------------------------

def post_process_docx(docx_path: Path) -> None:
    """
    Open the pandoc-generated docx and apply final styling:
    - Re-insert images via python-docx add_picture() (fixes LibreOffice rendering)
    - Colour header row of every table (navy fill, white bold text)
    - Centre all images / figures
    """
    import zipfile as _zf
    import re as _re
    from io import BytesIO
    from docx import Document
    from docx.shared import RGBColor, Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

    # ── 1. Read image bytes from the docx archive before opening Document ───
    media_map: dict[str, bytes] = {}
    with _zf.ZipFile(str(docx_path), "r") as z:
        rels_xml = z.read("word/_rels/document.xml.rels").decode()
        for m in _re.finditer(r"<Relationship[^>]*/>", rels_xml):
            entry = m.group(0)
            if "/image" not in entry:
                continue
            rid_m = _re.search(r'Id="(rId\d+)"', entry)
            tgt_m = _re.search(r'Target="(media/[^"]+)"', entry)
            if rid_m and tgt_m:
                media_map[rid_m.group(1)] = z.read(f"word/{tgt_m.group(1)}")

    doc = Document(str(docx_path))

    # ── 2. Re-insert each image using python-docx add_picture() ─────────────
    # Pandoc emits complex OOXML with non-standard attributes that LibreOffice
    # can't render.  python-docx produces minimal, standards-compliant markup.
    for para in doc.paragraphs:
        blips = para._p.findall(".//" + qn("a:blip"))
        if not blips:
            continue
        old_rid = blips[0].get(qn("r:embed"))
        img_bytes = media_map.get(old_rid)
        if img_bytes is None:
            continue

        # Preserve paragraph properties (indentation, style, etc.)
        pPr = para._p.find(qn("w:pPr"))
        for child in list(para._p):
            if child is not pPr:
                para._p.remove(child)

        # Insert fresh image run — simple, LibreOffice-compatible OOXML
        run = para.add_run()
        run.add_picture(BytesIO(img_bytes), width=Inches(6))
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after  = Pt(8)

    # ── 3. Style tables ──────────────────────────────────────────────────────
    for table in doc.tables:
        tbl_style = _get_style(doc, "Table Grid")
        if tbl_style:
            table.style = tbl_style
        for i, row in enumerate(table.rows):
            for cell in row.cells:
                if i == 0:
                    # Header row: navy background, white bold text
                    _set_cell_shading(cell, NAVY)
                    for para in cell.paragraphs:
                        for run in para.runs:
                            r = (WHITE >> 16) & 0xFF
                            g = (WHITE >> 8)  & 0xFF
                            b =  WHITE        & 0xFF
                            run.font.color.rgb = RGBColor(r, g, b)
                            run.font.bold = True
                            run.font.name = "Calibri"
                            run.font.size = Pt(10)
                elif i % 2 == 0:
                    # Alternate rows: light blue tint
                    _set_cell_shading(cell, ALT)
                # All cells: Calibri 10pt
                for para in cell.paragraphs:
                    para.paragraph_format.space_before = Pt(2)
                    para.paragraph_format.space_after  = Pt(2)
                    for run in para.runs:
                        run.font.name = "Calibri"
                        if run.font.size is None:
                            run.font.size = Pt(10)

    doc.save(str(docx_path))
    print(f"  Post-processing complete → {docx_path}")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> int:
    parser = argparse.ArgumentParser(description="Build Word report from Markdown")
    parser.add_argument("--md", default=None,
                        help="Path to markdown file (default: latest in reports/)")
    args = parser.parse_args()

    reports_dir = ROOT / "reports"

    # Resolve markdown file
    if args.md:
        md_path = Path(args.md).resolve()
    else:
        candidates = sorted(reports_dir.glob("validation_*.md"),
                            key=lambda p: p.stat().st_mtime, reverse=True)
        if not candidates:
            print("ERROR: no validation_*.md found in reports/", file=sys.stderr)
            return 1
        md_path = candidates[0]

    out_path = md_path.with_suffix(".docx")
    ref_base = reports_dir / "reference_base.docx"
    ref_styled = reports_dir / "reference_styled.docx"

    if not ref_base.exists():
        print("Extracting pandoc default reference.docx …")
        subprocess.run(
            ["pandoc", "--print-default-data-file", "reference.docx"],
            stdout=ref_base.open("wb"), check=True
        )

    print("Applying professional styles to reference template …")
    build_reference_docx(ref_base, ref_styled)

    print(f"Running pandoc: {md_path.name} → {out_path.name} …")

    # Rewrite all relative image paths to absolute so pandoc resolves them
    # regardless of table-adjacent context (pandoc resource-path has edge cases
    # with images that immediately follow large tables).
    import re
    md_text = md_path.read_text(encoding="utf-8")
    def _abs_img(m):
        alt, rel = m.group(1), m.group(2)
        abs_p = (reports_dir / rel).resolve()
        return f"![{alt}]({abs_p})"
    md_abs = re.sub(r'!\[([^\]]*)\]\(([^)]+\.png)\)', _abs_img, md_text)
    # Replace standalone '---' horizontal rules with '* * *'.
    # Pandoc's pipe_tables extension misparses '---' after a pipe table in
    # certain contexts and silently drops images that appear earlier in the
    # same section.  '* * *' is an unambiguous thematic break.
    md_abs = re.sub(r'^---$', '* * *', md_abs, flags=re.MULTILINE)
    md_tmp = md_path.with_name("_tmp_build.md")
    md_tmp.write_text(md_abs, encoding="utf-8")

    cmd = [
        "pandoc",
        str(md_tmp),
        "--from", "markdown-yaml_metadata_block+tex_math_dollars+pipe_tables+raw_html",
        "--to", "docx",
        "--output", str(out_path),
        "--reference-doc", str(ref_styled),
        "--highlight-style", "tango",
        "--lua-filter", str(ROOT / "scripts" / "docx_image_scale.lua"),
        "--standalone",
    ]

    # Lua filter may not exist yet — build it inline
    lua_filter = ROOT / "scripts" / "docx_image_scale.lua"
    if not lua_filter.exists():
        lua_filter.write_text(
            '-- Scale embedded images to fit page width\n'
            'function Image(el)\n'
            '  el.attributes["width"] = "6in"\n'
            '  return el\n'
            'end\n'
        )

    result = subprocess.run(cmd, capture_output=True, text=True)
    md_tmp.unlink(missing_ok=True)   # clean up temp file
    if result.returncode != 0:
        print("pandoc stderr:", result.stderr, file=sys.stderr)
        return result.returncode

    print("Post-processing table colours and image alignment …")
    post_process_docx(out_path)

    print(f"\nDone → {out_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
